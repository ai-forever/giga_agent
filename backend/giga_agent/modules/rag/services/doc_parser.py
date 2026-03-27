import io

from docx import Document as DocxDocument
from langchain_core.document_loaders.base import BaseBlobParser
from langchain_core.documents.base import Blob, Document


class CustomDocxParser(BaseBlobParser):
    """Parse DOCX files using python-docx library without unstructured dependency."""

    def lazy_parse(self, blob: Blob):
        """Parse DOCX file and yield a single Document."""
        docx_file = io.BytesIO(blob.as_bytes())
        doc = DocxDocument(docx_file)

        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        metadata = {
            "source": blob.source or "docx",
        }

        yield Document(page_content=full_text, metadata=metadata)