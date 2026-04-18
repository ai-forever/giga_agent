from __future__ import annotations

import asyncio
import io
import json
import uuid
from typing import Annotated
from urllib.request import urlopen

from docx import Document as DocxDocument
from langchain.tools import ToolRuntime, tool

from giga_agent.core.db import get_session_factory
from giga_agent.core.logging import get_logger
from giga_agent.embeddings.manager import EmbeddingManager
from giga_agent.modules.rag.database.collection_names import (
    rag_qdrant_collection_name_for_embedding,
)
from giga_agent.vectorstores.qdrant import (
    get_qdrant_client,
    resolve_qdrant_collection,
)
from giga_agent.modules.rag.database.qdrant_store import build_filter, search_chunks
from giga_agent.models.rag import RagCollectionsRepository
from giga_agent.core.agent.types import Collection

logger = get_logger(__name__)

MAX_READ_FILE_CHARS = 100_000
DEFAULT_READ_FILE_LIMIT = 1000


def _extract_pdf_text(data: bytes) -> str | None:
    from pdfminer.high_level import extract_text

    try:
        return extract_text(io.BytesIO(data))
    except Exception:
        return None


def _extract_docx_text(data: bytes) -> str | None:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception:
        return None

    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _decode_text_bytes(data: bytes) -> str | None:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return None


def _text_from_file_bytes(
    *,
    data: bytes,
    sandbox_path: str,
    file_name: str | None = None,
    media_type: str | None = None,
) -> str | None:
    normalized_media_type = (media_type or "").split(";", 1)[0].strip().lower()
    lower_name = (file_name or "").lower()
    lower_path = sandbox_path.lower()

    if (
        normalized_media_type == "application/pdf"
        or lower_name.endswith(".pdf")
        or lower_path.endswith(".pdf")
        or data.startswith(b"%PDF-")
    ):
        return _extract_pdf_text(data)

    if (
        normalized_media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower_name.endswith(".docx")
        or lower_path.endswith(".docx")
    ):
        return _extract_docx_text(data)

    return _decode_text_bytes(data)


async def _download_redirect_bytes(url: str) -> bytes:
    def _read_bytes() -> bytes:
        with urlopen(url, timeout=30.0) as response:
            return response.read()

    return await asyncio.to_thread(_read_bytes)


def _slice_lines(
    *,
    lines: list[str],
    offset: int | None,
    limit: int,
) -> tuple[list[str], int, int]:
    total_lines = len(lines)
    if offset is None:
        start_index = 0
    elif offset > 0:
        start_index = max(offset - 1, 0)
    else:
        start_index = max(total_lines + offset, 0)

    if limit <= 0:
        raise ValueError("limit must be greater than 0")
    end_index = min(start_index + limit, total_lines)

    return lines[start_index:end_index], start_index, end_index


def _format_numbered_lines(
    *,
    lines: list[str],
    start_line_number: int,
    max_chars: int = MAX_READ_FILE_CHARS,
) -> tuple[str, int, bool]:
    parts: list[str] = []
    used_chars = 0
    returned_lines = 0

    for index, line in enumerate(lines, start=start_line_number):
        prefix = "" if not parts else "\n"
        formatted_line = f"{index}|{line}"
        budget = max_chars - used_chars - len(prefix)
        if budget <= 0:
            return "".join(parts), returned_lines, True

        if len(formatted_line) > budget:
            if not parts:
                parts.append(prefix + formatted_line[:budget])
                returned_lines = 1
            return "".join(parts), returned_lines, True

        parts.append(prefix + formatted_line)
        used_chars += len(prefix) + len(formatted_line)
        returned_lines += 1

    return "".join(parts), returned_lines, False


def _build_next_read_hint(
    *,
    sandbox_path: str,
    total_lines: int,
    next_offset: int | None,
    remaining_lines: int,
    requested_limit: int | None,
    returned_lines: int,
    truncated_by_chars: bool,
) -> str:
    if total_lines == 0:
        return (
            "File is empty. Если хочешь перепроверить, вызови read_file с тем же "
            f"sandbox_path={sandbox_path!r}."
        )

    if remaining_lines <= 0 or next_offset is None:
        return (
            "Достигнут конец файла. Если нужно перечитать фрагмент, используй "
            f"read_file(sandbox_path={sandbox_path!r}, offset=1)."
        )

    next_limit = requested_limit or DEFAULT_READ_FILE_LIMIT
    char_limit_hint = ""
    if truncated_by_chars and returned_lines < next_limit:
        char_limit_hint = (
            f" Чтение файла ограничено лимитом в {MAX_READ_FILE_CHARS} символов, "
            "поэтому вернулось меньше строк, чем ты запрашивал."
        )

    return (
        f"Файл еще имеет {remaining_lines} строк после текущего фрагмента. "
        f"{char_limit_hint}"
        f"Чтобы продолжить, лучше вызвать "
        f"read_file(sandbox_path={sandbox_path!r}, offset={next_offset}, limit={next_limit})."
    )


@tool(
    description="""Семантический поиск по базе знаний пользователя через векторный поиск.

Используй для поиска информации из документов пользователя. Формулируй query как естественный вопрос.
При недостатке информации делай повторные запросы с другими формулировками.
Всегда цитируй источники (ID документа) в ответах.""",
)
async def get_documents(
    collection_uuid: Annotated[str, "UUID-коллекции"],
    query: Annotated[str, "Поисковый запрос для поиска релевантных документов"],
    runtime: ToolRuntime,
    limit: Annotated[int, "Количество документов, которые возвращаются"] = 10,
) -> dict:
    def _json(payload: dict) -> str:
        return json.dumps(payload, ensure_ascii=False, separators=(",", ":"))

    hint = (
        "Чтобы изучить документ подробнее, используй инструмент read_file с sandbox_path из результата. "
        "Это позволит прочитать файл целиком и найти нужную информацию."
    )

    if runtime is None:
        return {"error": "ToolRuntime is required", "documents": [], "hint": hint}

    user_id = runtime.config["configurable"]["langgraph_auth_user"]["identity"]
    owner_id = uuid.UUID(user_id) if isinstance(user_id, str) else user_id

    try:
        collection_id = uuid.UUID(collection_uuid)
    except ValueError:
        return {"error": "Invalid collection UUID", "documents": [], "hint": hint}

    factory = await get_session_factory()
    async with factory() as session:
        collection = await RagCollectionsRepository(session).get_by_id(
            owner_id=owner_id,
            collection_id=collection_id,
        )
        if collection is None:
            return _json({"error": "Collection not found", "documents": [], "hint": hint})

        embedding_runtime = await EmbeddingManager.resolve_by_id(
            collection.embedding_id,
            session=session,
        )
        embeddings = await embedding_runtime.get_embeddings()

    qdrant_client = get_qdrant_client()
    try:
        qfilter = build_filter(owner_id=owner_id, collection_id=collection_id)
        query_vector = (
            await embeddings.aembed_query(query)
            if hasattr(embeddings, "aembed_query")
            else await asyncio.to_thread(embeddings.embed_query, query)
        )
        qdrant_collection = await resolve_qdrant_collection(
            client=qdrant_client,
            collection_name=rag_qdrant_collection_name_for_embedding(collection.embedding_id),
            vector_size=len(query_vector),
        )
        points = await search_chunks(
            client=qdrant_client,
            collection_name=qdrant_collection,
            query_vector=query_vector,
            query_filter=qfilter,
            limit=limit,
        )
    except Exception as e:
        return {"error": str(e), "documents": [], "hint": hint}

    documents: list[dict] = []
    for p in points:
        payload = p.payload or {}
        chunk_id = str(p.id)
        doc_id = payload.get("document_id") or payload.get("file_id")
        name = payload.get("document_name") or payload.get("name")
        start_index = payload.get("start_index")
        end_index = payload.get("end_index")
        sandbox_path = payload.get("sandbox_path")
        content = payload.get("page_content")

        documents.append(
            {
                "chunk_id": chunk_id,
                "document_id": str(doc_id) if doc_id else None,
                "name": str(name) if name else None,
                "start_index": start_index,
                "end_index": end_index,
                "sandbox_path": str(sandbox_path) if sandbox_path else None,
                "content": str(content) if content else "",
                "score": getattr(p, "score", None),
            }
        )

    return {
        "collection_uuid": str(collection_id),
        "query": query,
        "limit": limit,
        "documents": documents,
        "hint": hint,
        "next_step": "Если информации недостаточно, используй read_file(sandbox_path) чтобы прочитать документ целиком, или переформулируй запрос.",
    }


@tool(
    description="""Читает файл из файловой системы. Ты можешь получить доступ к любому файлу напрямую с помощью этого инструмента.
Если пользователь передаёт путь к файлу, считай этот путь валидным. Если файл не существует, будет возвращена ошибка.
Использование:
- Можно опционально указать offset и limit по строкам, что особенно удобно для длинных файлов, но по умолчанию читаются первые 1000 строк файла. Чтобы читать его дальше повышай offset и читай файл дальше.
- Строки в результате нумеруются начиная с 1 в формате: НОМЕР_СТРОКИ|СОДЕРЖИМОЕ_СТРОКИ
- Если файл существует, но пустой, инструмент вернёт 'File is empty.'
- PDF и DOCX-файлы автоматически конвертируются в текст (с теми же ограничениями на размер ответа, что и для остальных файлов).""",
    extras={"repl_skip": True, "not_compress": True},
)
async def read_file(
    sandbox_path: Annotated[str, "Путь к файлу в sandbox (sandbox_path из результата get_documents)"],
    runtime: ToolRuntime,
    offset: Annotated[
        int,
        "Номер строки для начала чтения: положительные значения 1-indexed, отрицательные считаются от конца файла",
    ] = 1,
    limit: Annotated[
        int,
        "Количество строк для чтения. Если не передано, инструмент пытается вернуть файл целиком в пределах лимита размера ответа",
    ] = 1000,
) -> dict:
    """Читает файл из sandbox построчно с поддержкой offset/limit."""
    from giga_agent.sandbox.base import ContentResult, RedirectResult, StreamResult
    from giga_agent.sandbox.manager import SandboxManager

    if runtime is None:
        return {"error": "ToolRuntime is required", "content": None}

    user_id = runtime.config["configurable"]["langgraph_auth_user"]["identity"]
    owner_id = uuid.UUID(user_id) if isinstance(user_id, str) else user_id

    factory = await get_session_factory()
    async with factory() as session:
        manager = SandboxManager(session)
        try:
            file_record, result = await manager.read_file_by_path_for_user(
                user_id=owner_id,
                sandbox_path=sandbox_path,
            )
        except Exception as e:
            logger.warning("read_file failed for %s: %s", sandbox_path, e)
            return {"error": str(e), "content": None}

    text: str | None
    if isinstance(result, ContentResult):
        text = _text_from_file_bytes(
            data=result.data,
            sandbox_path=sandbox_path,
            file_name=file_record.original_name,
            media_type=result.media_type,
        )
        if text is None:
            return {
                "error": "Файл не является текстовым (бинарный формат)",
                "file": file_record.original_name,
                "size": file_record.size,
                "content": None,
            }
    elif isinstance(result, StreamResult):
        chunks = []
        async for chunk in result.stream:
            chunks.append(chunk)
        data = b"".join(chunks)
        text = _text_from_file_bytes(
            data=data,
            sandbox_path=sandbox_path,
            file_name=file_record.original_name,
            media_type=result.media_type,
        )
        if text is None:
            return {
                "error": "Файл не является текстовым (бинарный формат)",
                "file": file_record.original_name,
                "size": file_record.size,
                "content": None,
            }
    elif isinstance(result, RedirectResult):
        try:
            data = await _download_redirect_bytes(result.url)
        except Exception as e:
            logger.warning("read_file redirect download failed for %s: %s", sandbox_path, e)
            return {"error": str(e), "content": None}
        text = _text_from_file_bytes(
            data=data,
            sandbox_path=sandbox_path,
            file_name=file_record.original_name,
        )
        if text is None:
            return {
                "error": "Файл не является текстовым (бинарный формат)",
                "file": file_record.original_name,
                "size": file_record.size,
                "content": None,
            }
    else:
        return {"error": "Файл доступен только по URL, прямое чтение невозможно", "content": None}

    if text == "":
        return {
            "file": file_record.original_name,
            "sandbox_path": sandbox_path,
            "size": file_record.size,
            "offset": offset,
            "limit": limit,
            "total_lines": 0,
            "returned_lines": 0,
            "remaining_lines": 0,
            "truncated": False,
            "content": "File is empty.",
            "next_read_hint": _build_next_read_hint(
                sandbox_path=sandbox_path,
                total_lines=0,
                next_offset=None,
                remaining_lines=0,
                requested_limit=limit,
                returned_lines=0,
                truncated_by_chars=False,
            ),
        }

    lines = text.splitlines()
    try:
        selected_lines, start_index, _ = _slice_lines(
            lines=lines,
            offset=offset,
            limit=limit,
        )
    except ValueError as e:
        return {
            "error": str(e),
            "file": file_record.original_name,
            "sandbox_path": sandbox_path,
            "content": None,
        }

    content, returned_lines, truncated_by_chars = _format_numbered_lines(
        lines=selected_lines,
        start_line_number=start_index + 1,
    )
    remaining_lines = max(len(lines) - (start_index + returned_lines), 0)
    next_offset = start_index + returned_lines + 1 if remaining_lines > 0 else None
    truncated = truncated_by_chars or remaining_lines > 0

    return {
        "file": file_record.original_name,
        "sandbox_path": sandbox_path,
        "size": file_record.size,
        "offset": offset,
        "limit": limit,
        "total_lines": len(lines),
        "returned_lines": returned_lines,
        "remaining_lines": remaining_lines,
        "truncated": truncated,
        "content": content,
        "next_read_hint": _build_next_read_hint(
            sandbox_path=sandbox_path,
            total_lines=len(lines),
            next_offset=next_offset,
            remaining_lines=remaining_lines,
            requested_limit=limit,
            returned_lines=returned_lines,
            truncated_by_chars=truncated_by_chars,
        ),
    }


def has_collections(state):
    return len(state.get("collections", [])) > 0


RAG_PROMPT = """
====
БАЗА ЗНАНИЙ

У тебя есть доступ к документам пользователя через инструменты:
- get_documents — семантический (векторный) поиск по чанкам документов
- read_file — чтение файла целиком по sandbox_path

ВСЕГДА проверяй информацию в базе знаний перед ответом, даже если уверен в своих знаниях.

ДОСТУПНЫЕ КОЛЛЕКЦИИ:
{0}

СТРАТЕГИЯ РАБОТЫ:

1. ПРОСТЫЕ ЗАПРОСЫ (конкретный факт/процедура):
   * Сформулируй query как естественный вопрос с ключевыми терминами
   * Начни с limit=5-10
   * Если результат неполный → используй read_file(sandbox_path) чтобы прочитать документ целиком
   * Для смежных коллекций делай отдельные запросы

2. ГЛУБОКИЙ АНАЛИЗ:
   * Шаг 1: Обзорный запрос через get_documents -> определи структуру и ключевые термины
   * Шаг 2: Используй read_file для полного чтения ключевых документов
   * Шаг 3: Серия целевых запросов по каждому аспекту к базе знаний
   * Шаг 4: Структурированный отчет:
     - Резюме и выводы
     - Ключевые условия/риски с цитатами
     - Вопросы для уточнения
     - Таблица основных параметров

ВАЖНО: Если get_documents нашёл релевантный документ, но информации в чанке недостаточно — ОБЯЗАТЕЛЬНО прочитай файл целиком через read_file(sandbox_path), прежде чем говорить что данных нет.

ЦИТИРОВАНИЕ: Всегда указывай название документа и, если есть, раздел/пункт/страницу.

"""


def get_rag_info(collections: list[Collection]):
    if not collections:
        return RAG_PROMPT.format(
            "ДОСТУПНЫХ КОЛЛЕКЦИЙ НЕТ, ЕСЛИ ПОЛЬЗОВАТЕЛЬ "
            "ЗАПРАШИВАЕТ get_documents скажи,"
            " что ему нужно либо добавить документы, либо включить их в интерфейсе"
        )
    descriptions = []
    for collection in collections:
        description = (
            f"Название коллекции: {collection['name']}\nUUID: {collection['uuid']}"
        )
        if collection.get("metadata", {}).get("description"):
            description += (
                f"\nОписание коллекции: {collection['metadata']['description']}"
            )
        descriptions.append(description)
    return RAG_PROMPT.format("\n---\n".join(descriptions))
