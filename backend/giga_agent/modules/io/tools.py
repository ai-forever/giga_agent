from __future__ import annotations

import io
import mimetypes
import re
import uuid
from pathlib import PurePosixPath
from typing import Annotated

from docx import Document as DocxDocument
from langchain.tools import ToolRuntime, tool
from langchain_core.messages import AIMessage, ToolMessage
from langgraph.types import Command

from giga_agent.conf import get_settings
from giga_agent.core.agent.tool_policy import (
    ToolConfirmation,
    ToolEffect,
    tool_extras,
)
from giga_agent.core.db import get_session_factory
from giga_agent.core.logging import get_logger
from giga_agent.modules.io.memory_bridge import (
    _is_memory_path,
    _memory_delete,
    _memory_edit,
    _memory_read,
    _memory_write,
)
from giga_agent.sandbox.materialize import (
    materialize_bounded as _materialize_bounded,
)
from giga_agent.sandbox.materialize import (
    metadata_size as _metadata_size,
)
from giga_agent.utils.langgraph_sdk import get_user_id_from_config

logger = get_logger(__name__)

MAX_READ_FILE_CHARS = 100_000
DEFAULT_READ_FILE_LIMIT = 1000
# read_file/edit_file/write_file загружают файл целиком в память, поэтому работать
# с большими файлами через них запрещаем — для них есть shell (sed/tail/grep) и python.
# Один общий порог: текстовых файлов крупнее в этих сценариях практически не бывает.
MAX_TEXT_FILE_BYTES = 5 * 1024 * 1024
TABULAR_FILE_EXTENSIONS = frozenset(
    {
        ".csv",
        ".csv.gz",
        ".tsv",
        ".tsv.gz",
        ".xls",
        ".xlsx",
        ".xlsm",
        ".xlsb",
        ".ods",
        ".parquet",
        ".pq",
        ".feather",
        ".arrow",
        ".orc",
    }
)
TABULAR_MEDIA_TYPES = frozenset(
    {
        "text/csv",
        "text/tab-separated-values",
        "application/vnd.ms-excel",
        "application/vnd.ms-excel.sheet.binary.macroenabled.12",
        "application/vnd.ms-excel.sheet.macroenabled.12",
        "application/vnd.oasis.opendocument.spreadsheet",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.template",
        "application/vnd.apache.parquet",
        "application/x-parquet",
        "application/vnd.apache.arrow.file",
        "application/x-feather",
        "application/vnd.apache.orc",
    }
)


def _normalize_media_type(media_type: str | None) -> str:
    return (media_type or "").split(";", 1)[0].strip().lower()


def _is_tabular_file_reference(
    *,
    sandbox_path: str,
    file_name: str | None = None,
    media_type: str | None = None,
) -> bool:
    normalized_media_type = _normalize_media_type(media_type)
    if normalized_media_type in TABULAR_MEDIA_TYPES:
        return True

    for candidate in (sandbox_path, file_name):
        lower_candidate = (candidate or "").lower()
        if lower_candidate.endswith(tuple(TABULAR_FILE_EXTENSIONS)):
            return True

        guessed_media_type, _ = mimetypes.guess_type(lower_candidate)
        if _normalize_media_type(guessed_media_type) in TABULAR_MEDIA_TYPES:
            return True

    return False


def _build_tabular_read_hint(*, sandbox_path: str) -> str:
    return (
        f"Файл: {sandbox_path}\n"
        "Это похоже на табличные данные. read_file не читает такие файлы напрямую. "
        "Используй python tool и читай файл через подходящую библиотеку "
        "(например pandas, csv, openpyxl или pyarrow)."
    )


def _extract_pdf_text(data: bytes) -> str | None:
    from pdfminer.high_level import extract_text

    try:
        return extract_text(io.BytesIO(data))
    except Exception:
        return None


def _extract_docx_text(data: bytes) -> str | None:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception:
        return None

    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


def _decode_text_bytes(data: bytes) -> str | None:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return None


def _text_from_file_bytes(
    *,
    data: bytes,
    sandbox_path: str,
    file_name: str | None = None,
    media_type: str | None = None,
) -> str | None:
    normalized_media_type = _normalize_media_type(media_type)
    lower_name = (file_name or "").lower()
    lower_path = sandbox_path.lower()

    if (
        normalized_media_type == "application/pdf"
        or lower_name.endswith(".pdf")
        or lower_path.endswith(".pdf")
        or data.startswith(b"%PDF-")
    ):
        return _extract_pdf_text(data)

    if (
        normalized_media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower_name.endswith(".docx")
        or lower_path.endswith(".docx")
    ):
        return _extract_docx_text(data)

    return _decode_text_bytes(data)


def _slice_lines(
    *,
    lines: list[str],
    offset: int | None,
    limit: int,
) -> tuple[list[str], int, int]:
    total_lines = len(lines)
    if offset is None:
        start_index = 0
    elif offset > 0:
        start_index = max(offset - 1, 0)
    else:
        start_index = max(total_lines + offset, 0)

    if limit <= 0:
        raise ValueError("limit must be greater than 0")
    end_index = min(start_index + limit, total_lines)

    return lines[start_index:end_index], start_index, end_index


def _format_numbered_lines(
    *,
    lines: list[str],
    start_line_number: int,
    max_chars: int = MAX_READ_FILE_CHARS,
) -> tuple[str, int, bool]:
    parts: list[str] = []
    used_chars = 0
    returned_lines = 0

    for index, line in enumerate(lines, start=start_line_number):
        prefix = "" if not parts else "\n"
        formatted_line = f"{index:>6}|{line}"
        budget = max_chars - used_chars - len(prefix)
        if budget <= 0:
            return "".join(parts), returned_lines, True

        if len(formatted_line) > budget:
            if not parts:
                parts.append(prefix + formatted_line[:budget])
                returned_lines = 1
            return "".join(parts), returned_lines, True

        parts.append(prefix + formatted_line)
        used_chars += len(prefix) + len(formatted_line)
        returned_lines += 1

    return "".join(parts), returned_lines, False


def _build_next_read_hint(
    *,
    sandbox_path: str,
    total_lines: int,
    next_offset: int | None,
    remaining_lines: int,
    requested_limit: int | None,
    returned_lines: int,
    truncated_by_chars: bool,
) -> str:
    if total_lines == 0:
        return (
            "File is empty. Если хочешь перепроверить, вызови read_file с тем же "
            f"sandbox_path={sandbox_path!r}."
        )

    if remaining_lines <= 0 or next_offset is None:
        return (
            "Достигнут конец файла. Если нужно перечитать фрагмент, используй "
            f"read_file(sandbox_path={sandbox_path!r}, offset=1)."
        )

    next_limit = requested_limit or DEFAULT_READ_FILE_LIMIT
    char_limit_hint = ""
    if truncated_by_chars and returned_lines < next_limit:
        char_limit_hint = (
            f" Чтение файла ограничено лимитом в {MAX_READ_FILE_CHARS} символов, "
            "поэтому вернулось меньше строк, чем ты запрашивал."
        )

    return (
        f"Файл еще имеет {remaining_lines} строк после текущего фрагмента. "
        f"{char_limit_hint}"
        f"Чтобы продолжить, лучше вызвать "
        f"read_file(sandbox_path={sandbox_path!r}, offset={next_offset}, limit={next_limit})."
    )


async def _get_owner_id(runtime: ToolRuntime) -> uuid.UUID:
    user_id = get_user_id_from_config(runtime.config)
    return uuid.UUID(user_id) if isinstance(user_id, str) else user_id


def _is_cli_runtime() -> bool:
    return get_settings().giga_agent_runtime == "cli"


async def _get_cli_sandbox_runtime(runtime: ToolRuntime):
    """Resolve the sandbox runtime in CLI mode without hitting the DB.

    Falls back to creating a resolver if the tool node hasn't injected one
    into the config yet (e.g. very early test harness).
    """
    from giga_agent.core.agent.runtime_resolver import RuntimeResolver
    from giga_agent.sandbox.manager.runtime_factory import SandboxRuntimeFactory

    try:
        resolver = RuntimeResolver.from_config(runtime.config)
    except ValueError:
        resolver = await RuntimeResolver.create(runtime.config)
    resolved = await resolver.get_sandbox()
    return SandboxRuntimeFactory.build(resolved.provider, resolved.sandbox)


class _CliFileStub:
    """Tiny shim mimicking the bits of `File` that read_file uses for hints."""

    __slots__ = ("sandbox_path", "original_name", "size")

    def __init__(self, sandbox_path: str) -> None:
        self.sandbox_path = sandbox_path
        self.original_name = PurePosixPath(sandbox_path).name or sandbox_path
        self.size = 0


def _looks_like_missing_file_error(error: Exception) -> bool:
    if isinstance(error, FileNotFoundError):
        return True
    message = str(error).lower()
    return "not found" in message or "не найден" in message or "no such file" in message


async def _find_relative_path_alternative(
    owner_id: uuid.UUID,
    absolute_path: str,
    runtime: ToolRuntime | None = None,
) -> str | None:
    """Если по абсолютному пути файла нет, проверяем относительные варианты
    (без ведущего слэша и просто basename). Возвращаем первый существующий
    относительный путь либо None."""
    if not absolute_path.startswith("/") or absolute_path.startswith("/bucket/"):
        return None

    candidates: list[str] = []
    stripped = absolute_path.lstrip("/")
    if stripped and "/" in stripped:
        candidates.append(stripped)
    basename = PurePosixPath(absolute_path).name
    if basename and basename not in candidates:
        candidates.append(basename)

    if not candidates:
        return None

    if _is_cli_runtime() and runtime is not None:
        try:
            sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
        except Exception:
            return None
        for candidate in candidates:
            try:
                exists = await sandbox_runtime.file_exists(candidate)
            except Exception:
                continue
            if exists:
                return candidate
        return None

    from giga_agent.sandbox.manager import SandboxManager

    factory = await get_session_factory()
    async with factory() as session:
        manager = SandboxManager(session)
        for candidate in candidates:
            try:
                exists = await manager.file_exists_for_user(
                    user_id=owner_id,
                    sandbox_path=candidate,
                )
            except Exception:
                continue
            if exists:
                return candidate
    return None


def _format_relative_path_hint(absolute_path: str, relative_path: str) -> str:
    return (
        f" Подсказка: по абсолютному пути {absolute_path!r} файла нет, "
        f'но похожий файл существует по относительному пути "{relative_path!r}", '
        f"возможно ты имел в виду его."
    )


async def _get_current_workdir(
    owner_id: uuid.UUID,
    runtime: ToolRuntime | None = None,
) -> str | None:
    """Дешёво достаёт cwd рантайма пользователя. Возвращает None, если не
    удалось (например, рантайм не поддерживает cwd или sandbox не настроен)."""
    if _is_cli_runtime() and runtime is not None:
        try:
            sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
        except Exception:
            return None
        try:
            return sandbox_runtime.current_workdir()
        except Exception:
            return None

    from giga_agent.sandbox.manager import SandboxManager

    try:
        factory = await get_session_factory()
        async with factory() as session:
            manager = SandboxManager(session)
            return await manager.get_current_workdir_for_user(user_id=owner_id)
    except Exception:
        return None


async def _build_missing_file_hint(
    owner_id: uuid.UUID,
    absolute_path: str,
    runtime: ToolRuntime | None = None,
) -> str:
    """Строит подсказку для случая «файл не найден»: текущая рабочая
    директория и, если есть, относительный путь, по которому файл существует."""
    cwd = await _get_current_workdir(owner_id, runtime)
    alt = await _find_relative_path_alternative(owner_id, absolute_path, runtime)
    parts: list[str] = []
    if cwd:
        parts.append(f" Текущая рабочая директория: {cwd}.")
    if alt:
        parts.append(_format_relative_path_hint(absolute_path, alt))
    return "".join(parts)


def _build_io_command(
    *,
    runtime: ToolRuntime,
    tool_name: str,
    content: str,
    is_error: bool = False,
) -> Command:
    tool_message_kwargs: dict = {
        "tool_call_id": getattr(runtime, "tool_call_id", "") or "",
        "content": content,
        "name": tool_name,
        "additional_kwargs": {"tool_name": tool_name},
    }
    if is_error:
        tool_message_kwargs["status"] = "error"
    return Command(update={"messages": [ToolMessage(**tool_message_kwargs)]})


def _overwrite_too_large_error(file_path: str, size: int | None) -> str:
    size_part = (
        f"{size} байт" if size is not None else f"больше {MAX_TEXT_FILE_BYTES} байт"
    )
    return (
        f"Ошибка: Файл {file_path} слишком большой для перезаписи "
        f"({size_part}, лимит {MAX_TEXT_FILE_BYTES}). "
        "Перезапиши его через shell-команду."
    )


def _edit_too_large_error(file_path: str, size: int | None) -> str:
    size_part = (
        f"{size} байт" if size is not None else f"больше {MAX_TEXT_FILE_BYTES} байт"
    )
    return (
        f"Ошибка: Файл {file_path} слишком большой для edit_file "
        f"({size_part}, лимит {MAX_TEXT_FILE_BYTES}). "
        "Редактируй его через shell-команду (например sed или python)."
    )


def _read_too_large_error(sandbox_path: str, size: int | None) -> str:
    size_part = (
        f"{size} байт" if size is not None else f"больше {MAX_TEXT_FILE_BYTES} байт"
    )
    return (
        f"Ошибка: Файл {sandbox_path} слишком большой для read_file "
        f"({size_part}, лимит {MAX_TEXT_FILE_BYTES}). "
        "Читай его через shell-команду (например sed -n, head, tail или grep) "
        "или через python tool."
    )


async def _probe_existing_file(
    file_path: str,
    owner_id: uuid.UUID,
    runtime: ToolRuntime | None = None,
):
    """Читает файл из рантайма для проверок перед записью/правкой.

    Возвращает (result, known_size). Для больших файлов result — ленивый
    StreamResult (поток ещё НЕ вычитан), known_size — размер из метаданных БД
    (или None в CLI). Это позволяет проверить размер, не материализуя гигабайты.
    """
    if _is_cli_runtime() and runtime is not None:
        sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
        result = await sandbox_runtime.read_file(file_path)
        return result, None

    from giga_agent.sandbox.manager import SandboxManager

    factory = await get_session_factory()
    async with factory() as session:
        manager = SandboxManager(session)
        file_record, result = await manager.read_file_by_path_for_user(
            user_id=owner_id,
            sandbox_path=file_path,
        )
    return result, getattr(file_record, "size", None)


async def _read_text_for_edit(
    file_path: str,
    owner_id: uuid.UUID,
    runtime: ToolRuntime | None = None,
) -> tuple[str | None, str | None]:
    """Читает текстовый файл для edit_file одним проходом.

    Возвращает (text, error_message): ровно одно из двух не None. edit_file держит
    файл целиком в памяти, поэтому слишком большие файлы отклоняются по метаданным
    ДО материализации (а поток дополнительно читается с ранним обрывом на лимите).
    """
    from giga_agent.sandbox.base import ContentResult

    try:
        result, known_size = await _probe_existing_file(file_path, owner_id, runtime)
    except Exception as e:
        logger.warning("edit_file read failed for %s: %s", file_path, e)
        return None, f"Ошибка: {e}"

    size = _metadata_size(result, known_size)
    # ContentResult уже материализован рантаймом (< LARGE_FILE_THRESHOLD), он безопасен.
    # Для остального с неизвестным размером (поток/redirect) читать целиком нельзя.
    if (size is not None and size > MAX_TEXT_FILE_BYTES) or (
        size is None and not isinstance(result, ContentResult)
    ):
        return None, _edit_too_large_error(file_path, size)

    data, too_large = await _materialize_bounded(result, MAX_TEXT_FILE_BYTES)
    if too_large:
        return None, _edit_too_large_error(file_path, None)
    if data is None:
        return None, "Ошибка: Файл доступен только по URL, прямое чтение невозможно"

    try:
        return data.decode("utf-8"), None
    except UnicodeDecodeError:
        return None, (
            f"Ошибка: Файл не является текстовым (бинарный формат). Файл: {file_path}"
        )


async def _check_overwrite_allowed(
    file_path: str,
    owner_id: uuid.UUID,
    runtime: ToolRuntime | None = None,
) -> tuple[str | None, int | None]:
    """Проверяет, можно ли перезаписать существующий файл.

    Возвращает (error_message, old_size). Если error_message не None — перезапись
    запрещена (файл слишком большой, бинарный или недоступен для прямого чтения).
    Иначе old_size — размер старого содержимого в байтах.

    Важно: НЕ материализует большие файлы в память. Размер проверяется по
    метаданным (DB size / content_length), а поток читается с ранним обрывом на
    лимите, чтобы перезапись гигабайтного файла не уронила сервер по памяти.
    """
    from giga_agent.sandbox.base import RedirectResult

    try:
        result, known_size = await _probe_existing_file(file_path, owner_id, runtime)
    except Exception as e:
        logger.warning("write_file overwrite probe failed for %s: %s", file_path, e)
        return (f"Ошибка: не удалось прочитать файл для перезаписи: {e}", None)

    size = _metadata_size(result, known_size)
    if size is not None and size > MAX_TEXT_FILE_BYTES:
        return (_overwrite_too_large_error(file_path, size), None)

    if isinstance(result, RedirectResult):
        return (
            f"Ошибка: Файл {file_path} хранится во внешнем хранилище и недоступен для "
            "прямой перезаписи через write_file. Используй shell-команду.",
            None,
        )

    data, too_large = await _materialize_bounded(result, MAX_TEXT_FILE_BYTES)
    if too_large:
        return (_overwrite_too_large_error(file_path, None), None)
    if data is None:
        return (
            f"Ошибка: Файл {file_path} доступен только по URL, перезапись через "
            "write_file невозможна. Используй shell-команду.",
            None,
        )

    try:
        data.decode("utf-8")
    except UnicodeDecodeError:
        return (
            f"Ошибка: Файл {file_path} не является текстовым (бинарный формат), "
            "перезапись через write_file не поддерживается. Используй shell-команду.",
            None,
        )

    return (None, len(data))


def _write_success_message(file_path: str, new_size: int, old_size: int | None) -> str:
    if old_size is None:
        return f"Файл создан: {file_path}, размер: {new_size} байт"
    return f"Файл перезаписан: {file_path} (было {old_size} → стало {new_size} байт)"


@tool(
    description="""Читает файл из файловой системы. Ты можешь получить доступ к любому файлу напрямую с помощью этого инструмента.
Если пользователь передаёт путь к файлу, считай этот путь валидным. Если файл не существует, будет возвращена ошибка.
Использование:
- Можно опционально указать offset и limit по строкам, что особенно удобно для длинных файлов, но по умолчанию читаются первые 1000 строк файла. Чтобы читать его дальше повышай offset и читай файл дальше.
- Строки в результате нумеруются начиная с 1 в формате: НОМЕР_СТРОКИ|СОДЕРЖИМОЕ_СТРОКИ
- Если файл существует, но пустой, инструмент вернёт 'File is empty.'
- Табличные файлы (например CSV, TSV, Excel, ODS, Parquet, Arrow, Feather, ORC) нужно читать через python tool, а не через read_file.
- PDF и DOCX-файлы автоматически конвертируются в текст (с теми же ограничениями на размер ответа, что и для остальных файлов).""",
    extras=tool_extras(
        ToolEffect.READ,
        repl_skip=True,
        not_compress=True,
        not_process=True,
    ),
)
async def read_file(
    sandbox_path: Annotated[
        str, "Путь к файлу в sandbox (sandbox_path из результата get_documents)"
    ],
    runtime: ToolRuntime,
    offset: Annotated[
        int,
        "Номер строки для начала чтения: положительные значения 1-indexed, отрицательные считаются от конца файла",
    ] = 1,
    limit: Annotated[
        int,
        "Количество строк для чтения. Если не передано, инструмент пытается вернуть файл целиком в пределах лимита размера ответа",
    ] = 1000,
) -> Command:
    """Читает файл из sandbox построчно с поддержкой offset/limit."""
    from giga_agent.sandbox.manager import SandboxManager

    def _result(text: str, *, is_error: bool = False) -> Command:
        return _build_io_command(
            runtime=runtime,
            tool_name="read_file",
            content=text,
            is_error=is_error,
        )

    if runtime is None:
        return _result("Ошибка: ToolRuntime is required", is_error=True)

    if _is_memory_path(sandbox_path):
        return await _memory_read(sandbox_path, runtime)

    user_id = get_user_id_from_config(runtime.config)
    owner_id = uuid.UUID(user_id) if isinstance(user_id, str) else user_id

    if _is_tabular_file_reference(sandbox_path=sandbox_path):
        return _result(_build_tabular_read_hint(sandbox_path=sandbox_path))

    if _is_cli_runtime():
        try:
            sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
            result = await sandbox_runtime.read_file(sandbox_path)
        except Exception as e:
            logger.warning("read_file failed for %s: %s", sandbox_path, e)
            error_text = f"Ошибка: {e}"
            if _looks_like_missing_file_error(e):
                error_text += await _build_missing_file_hint(
                    owner_id, sandbox_path, runtime
                )
            return _result(error_text, is_error=True)
        file_record = _CliFileStub(sandbox_path)
        # В CLI нет надёжного размера из метаданных (stub.size == 0), поэтому
        # полагаемся на content_length из result и потоковый обрыв ниже.
        known_size: int | None = None
    else:
        factory = await get_session_factory()
        async with factory() as session:
            manager = SandboxManager(session)
            try:
                file_record, result = await manager.read_file_by_path_for_user(
                    user_id=owner_id,
                    sandbox_path=sandbox_path,
                )
            except Exception as e:
                logger.warning("read_file failed for %s: %s", sandbox_path, e)
                error_text = f"Ошибка: {e}"
                if _looks_like_missing_file_error(e):
                    error_text += await _build_missing_file_hint(owner_id, sandbox_path)
                return _result(error_text, is_error=True)
        known_size = getattr(file_record, "size", None)

    media_type = getattr(result, "media_type", None)

    # Табличные файлы отсекаем по метаданным ДО любого чтения содержимого.
    if _is_tabular_file_reference(
        sandbox_path=sandbox_path,
        file_name=file_record.original_name,
        media_type=media_type,
    ):
        return _result(_build_tabular_read_hint(sandbox_path=sandbox_path))

    # Размер по метаданным: отказываем большим файлам, не материализуя их.
    size = _metadata_size(result, known_size)
    if size is not None and size > MAX_TEXT_FILE_BYTES:
        return _result(_read_too_large_error(sandbox_path, size), is_error=True)

    # Читаем с ранним обрывом на лимите — даже если размер не был известен заранее
    # (CLI/redirect без content_length), файл не сможет переполнить память.
    try:
        data, too_large = await _materialize_bounded(result, MAX_TEXT_FILE_BYTES)
    except Exception as e:
        logger.warning("read_file materialize failed for %s: %s", sandbox_path, e)
        return _result(f"Ошибка: {e}", is_error=True)

    if too_large:
        return _result(_read_too_large_error(sandbox_path, None), is_error=True)
    if data is None:
        return _result(
            "Ошибка: Файл доступен только по URL, прямое чтение невозможно",
            is_error=True,
        )

    text: str | None = _text_from_file_bytes(
        data=data,
        sandbox_path=sandbox_path,
        file_name=file_record.original_name,
        media_type=media_type,
    )
    if text is None:
        return _result(
            f"Ошибка: Файл не является текстовым (бинарный формат). "
            f"Файл: {file_record.original_name}, размер: {file_record.size}",
            is_error=True,
        )

    if text == "":
        hint = _build_next_read_hint(
            sandbox_path=sandbox_path,
            total_lines=0,
            next_offset=None,
            remaining_lines=0,
            requested_limit=limit,
            returned_lines=0,
            truncated_by_chars=False,
        )
        return _result(f"Файл: {sandbox_path}\nFile is empty.\n{hint}")

    lines = text.splitlines()
    try:
        selected_lines, start_index, _ = _slice_lines(
            lines=lines,
            offset=offset,
            limit=limit,
        )
    except ValueError as e:
        return _result(f"Ошибка: {e}", is_error=True)

    content, returned_lines, truncated_by_chars = _format_numbered_lines(
        lines=selected_lines,
        start_line_number=start_index + 1,
    )
    remaining_lines = max(len(lines) - (start_index + returned_lines), 0)
    next_offset = start_index + returned_lines + 1 if remaining_lines > 0 else None

    return _result(
        f"Файл: {sandbox_path}\n----\n"
        + content
        + "\n----\n"
        + _build_next_read_hint(
            sandbox_path=sandbox_path,
            total_lines=len(lines),
            next_offset=next_offset,
            remaining_lines=remaining_lines,
            requested_limit=limit,
            returned_lines=returned_lines,
            truncated_by_chars=truncated_by_chars,
        )
    )


@tool(
    description="""Создаёт новый файл в файловой системе или полностью перезаписывает существующий.

Использование:
- По умолчанию write_file создаёт новый файл. Если файл по указанному пути уже существует, будет возвращена ошибка.
- Чтобы перезаписать существующий файл целиком, передай overwrite=True. Это заменяет ВСЁ содержимое файла, а не подстроку — для точечных правок используй edit_file.
- Предпочитай edit_file для точечных изменений; overwrite=True уместен, когда нужно переписать файл целиком.""",
    extras=tool_extras(
        ToolEffect.WRITE,
        repl_skip=True,
        not_compress=True,
        not_process=True,
    ),
)
async def write_file(
    file_path: Annotated[
        str,
        "Абсолютный путь для создания файла. Должен быть абсолютным, не относительным.",
    ],
    content: Annotated[
        str,
        "Текстовое содержимое для записи в файл. Этот параметр обязателен.",
    ],
    runtime: ToolRuntime,
    overwrite: Annotated[
        bool,
        "Если True, полностью перезаписывает файл, когда он уже существует, вместо ошибки. По умолчанию False.",
    ] = False,
) -> Command:
    """Создаёт новый файл по указанному пути или перезаписывает существующий."""
    from giga_agent.sandbox.manager import SandboxManager

    def _result(text: str, *, is_error: bool = False) -> Command:
        return _build_io_command(
            runtime=runtime,
            tool_name="write_file",
            content=text,
            is_error=is_error,
        )

    if runtime is None:
        return _result("Ошибка: ToolRuntime is required", is_error=True)

    if _is_memory_path(file_path):
        return await _memory_write(file_path, content, runtime, overwrite=overwrite)

    owner_id = await _get_owner_id(runtime)
    content_bytes = content.encode("utf-8")
    exists_error = (
        f"Ошибка: Файл уже существует ({file_path}). "
        "Для точечной правки используй edit_file. "
        "Для полной перезаписи вызови write_file с overwrite=True."
    )

    if _is_cli_runtime():
        try:
            sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
            exists = await sandbox_runtime.file_exists(file_path)
        except Exception as e:
            logger.warning(
                "write_file file_exists check failed for %s: %s", file_path, e
            )
            return _result(f"Ошибка: {e}", is_error=True)

        old_size: int | None = None
        if exists:
            if not overwrite:
                return _result(exists_error, is_error=True)
            guard_err, old_size = await _check_overwrite_allowed(
                file_path, owner_id, runtime
            )
            if guard_err:
                return _result(guard_err, is_error=True)

        try:
            await sandbox_runtime.write_file_content(file_path, content_bytes)
        except Exception as e:
            logger.warning("write_file failed for %s: %s", file_path, e)
            return _result(f"Ошибка: {e}", is_error=True)

        return _result(_write_success_message(file_path, len(content_bytes), old_size))

    factory = await get_session_factory()
    async with factory() as session:
        manager = SandboxManager(session)
        try:
            exists = await manager.file_exists_for_user(
                user_id=owner_id,
                sandbox_path=file_path,
            )
        except Exception as e:
            logger.warning(
                "write_file file_exists check failed for %s: %s", file_path, e
            )
            return _result(f"Ошибка: {e}", is_error=True)

        old_size = None
        if exists:
            if not overwrite:
                return _result(exists_error, is_error=True)
            guard_err, old_size = await _check_overwrite_allowed(
                file_path, owner_id, runtime
            )
            if guard_err:
                return _result(guard_err, is_error=True)

        try:
            await manager.write_file_content_for_user(
                user_id=owner_id,
                sandbox_path=file_path,
                content=content_bytes,
            )
        except Exception as e:
            logger.warning("write_file failed for %s: %s", file_path, e)
            return _result(f"Ошибка: {e}", is_error=True)

    return _result(_write_success_message(file_path, len(content_bytes), old_size))


_FILE_MUTATING_TOOLS = frozenset({"edit_file", "write_file", "shell"})


def _has_recent_read_file(messages: list, file_path: str, lookback: int = 2) -> bool:
    """Return True if a recent read_file call targeted *file_path*.

    Walks messages backwards, skipping the first AIMessage (the current
    edit_file call), and inspects up to *lookback* preceding AIMessages.
    If the nearest read_file invocation has ``sandbox_path`` equal to
    *file_path*, returns True.
    """
    skipped_first_ai = False
    ai_count = 0
    for msg in reversed(messages):
        if not isinstance(msg, AIMessage):
            continue
        if not skipped_first_ai:
            skipped_first_ai = True
            continue
        ai_count += 1
        if ai_count > lookback:
            break
        for tc in reversed(msg.tool_calls or []):
            name = tc.get("name", "")
            if name == "read_file":
                args = tc.get("args") or {}
                return args.get("sandbox_path", "") == file_path
    return False


@tool(
    description="""Выполняет точную замену строк в файлах.

Использование:
- Перед редактированием необходимо прочитать файл. Этот инструмент выдаст ошибку, если файл не существует.
- При редактировании сохраняй точные отступы (табы/пробелы) из вывода read_file. Никогда не включай префиксы номеров строк в find_string или replace_string.
- ВСЕГДА предпочитай редактирование существующих файлов созданию новых.""",
    extras=tool_extras(
        ToolEffect.WRITE,
        repl_skip=True,
        not_compress=True,
        not_process=True,
    ),
)
async def edit_file(
    file_path: Annotated[
        str,
        "Абсолютный путь к редактируемому файлу. Должен быть абсолютным, не относительным.",
    ],
    find_string: Annotated[
        str,
        "Точный текст для поиска и замены. Должен быть уникальным в файле, если replace_all не True.",
    ],
    replace_string: Annotated[
        str,
        "Текст, на который заменяется find_string. Должен отличаться от find_string.",
    ],
    runtime: ToolRuntime,
    replace_all: Annotated[
        bool,
        "Если True, заменяет все вхождения find_string. Если False (по умолчанию), find_string должен быть уникальным.",
    ] = False,
) -> Command:
    """Выполняет точную замену строк в указанном файле."""
    from giga_agent.sandbox.manager import SandboxManager

    def _result(text: str, *, is_error: bool = False) -> Command:
        return _build_io_command(
            runtime=runtime,
            tool_name="edit_file",
            content=text,
            is_error=is_error,
        )

    if runtime is None:
        return _result("Ошибка: ToolRuntime is required", is_error=True)

    if _is_memory_path(file_path):
        return await _memory_edit(
            file_path, find_string, replace_string, replace_all, runtime
        )

    owner_id = await _get_owner_id(runtime)
    cli_mode = _is_cli_runtime()
    cli_sandbox_runtime = None

    if cli_mode:
        try:
            cli_sandbox_runtime = await _get_cli_sandbox_runtime(runtime)
            exists = await cli_sandbox_runtime.file_exists(file_path)
        except Exception as e:
            logger.warning(
                "edit_file file_exists check failed for %s: %s", file_path, e
            )
            return _result(f"Ошибка: {e}", is_error=True)

        if not exists:
            error_text = (
                f"Ошибка: Файл не существует ({file_path}). "
                "Используй write_file для создания новых файлов."
            )
            error_text += await _build_missing_file_hint(owner_id, file_path, runtime)
            return _result(error_text, is_error=True)
    else:
        factory = await get_session_factory()
        async with factory() as session:
            manager = SandboxManager(session)
            try:
                exists = await manager.file_exists_for_user(
                    user_id=owner_id,
                    sandbox_path=file_path,
                )
            except Exception as e:
                logger.warning(
                    "edit_file file_exists check failed for %s: %s", file_path, e
                )
                return _result(f"Ошибка: {e}", is_error=True)

            if not exists:
                error_text = (
                    f"Ошибка: Файл не существует ({file_path}). "
                    "Используй write_file для создания новых файлов."
                )
                error_text += await _build_missing_file_hint(owner_id, file_path)
                return _result(error_text, is_error=True)

    if find_string == replace_string:
        return _result(
            f"Ошибка: find_string и replace_string совпадают. Укажи разные значения. Файл: {file_path}",
            is_error=True,
        )

    raw_text, read_error = await _read_text_for_edit(file_path, owner_id, runtime)
    if read_error is not None:
        return _result(read_error, is_error=True)
    assert raw_text is not None

    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    find_string = find_string.replace("\r\n", "\n").replace("\r", "\n")
    replace_string = replace_string.replace("\r\n", "\n").replace("\r", "\n")

    count = text.count(find_string)
    if count == 0 and find_string.endswith("\n"):
        stripped = find_string[:-1]
        if text.endswith(stripped):
            find_string = stripped
            count = text.count(find_string)
    if count == 0:
        line_num_hint = ""
        if re.search(r"^\s{0,6}\d+\|", find_string, re.MULTILINE):
            line_num_hint = (
                " ВНИМАНИЕ: похоже, ты включил префиксы номеров строк "
                "(вида «  123|») из вывода read_file в find_string. "
                "Эти префиксы НЕ являются частью файла — убери их."
            )

        messages = runtime.state.get("messages", []) if runtime.state else []
        if _has_recent_read_file(messages, file_path):
            return _result(
                f"Ошибка: find_string не найден в файле {file_path}. "
                "Ты уже читал файл — ОБЯЗАТЕЛЬНО вызови think и внимательно сверь "
                "find_string с актуальным содержимым файла посимвольно (пробелы, "
                "отступы, переносы строк), прежде чем повторять попытку."
                + line_num_hint,
                is_error=True,
            )
        return _result(
            f"Ошибка: find_string не найден в файле {file_path}. "
            "Перечитай файл через read_file и вызови think, чтобы внимательно сверить "
            "find_string с актуальным содержимым файла, прежде чем повторять попытку."
            + line_num_hint,
            is_error=True,
        )

    if not replace_all and count > 1:
        return _result(
            f"Ошибка: find_string не уникален, найдено {count} вхождений в {file_path}. "
            "Передай более длинную строку с контекстом для уникальной идентификации, "
            "или используй replace_all=True для замены всех вхождений.",
            is_error=True,
        )

    if replace_all:
        new_text = text.replace(find_string, replace_string)
        replacements = count
    else:
        new_text = text.replace(find_string, replace_string, 1)
        replacements = 1

    new_bytes = new_text.encode("utf-8")

    if cli_mode:
        try:
            assert cli_sandbox_runtime is not None
            await cli_sandbox_runtime.write_file_content(file_path, new_bytes)
        except Exception as e:
            logger.warning("edit_file write failed for %s: %s", file_path, e)
            return _result(f"Ошибка: {e}", is_error=True)
    else:
        factory = await get_session_factory()
        async with factory() as session:
            manager = SandboxManager(session)
            try:
                await manager.write_file_content_for_user(
                    user_id=owner_id,
                    sandbox_path=file_path,
                    content=new_bytes,
                )
            except Exception as e:
                logger.warning("edit_file write failed for %s: %s", file_path, e)
                return _result(f"Ошибка: {e}", is_error=True)

    return _result(f"Файл отредактирован: {file_path}, замен: {replacements}")


@tool(
    description="""Удаляет файл памяти.

Использование:
- Поддерживается только для путей под /memories/. Для удаления обычных файлов используй shell.
- Удаляет файл целиком, индекс семантического поиска тоже очищается.""",
    extras=tool_extras(
        ToolEffect.DESTRUCTIVE,
        confirmation=ToolConfirmation.ALWAYS,
        repl_skip=True,
        not_compress=True,
        not_process=True,
    ),
)
async def delete_file(
    file_path: Annotated[
        str,
        "Абсолютный путь к файлу памяти под /memories/. Должен быть абсолютным.",
    ],
    runtime: ToolRuntime,
) -> Command:
    """Удаляет файл памяти (или возвращает ошибку для не-memory путей)."""

    def _result(text: str, *, is_error: bool = False) -> Command:
        return _build_io_command(
            runtime=runtime,
            tool_name="delete_file",
            content=text,
            is_error=is_error,
        )

    if runtime is None:
        return _result("Ошибка: ToolRuntime is required", is_error=True)

    if _is_memory_path(file_path):
        return await _memory_delete(file_path, runtime)

    return _result(
        "Ошибка: delete_file поддерживается только для путей под /memories/. "
        "Для удаления обычных файлов используй shell-команду.",
        is_error=True,
    )
